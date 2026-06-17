"""
Markdown Resume → Docx → Pages 构建脚本（技能模板）

使用方式：
  1. 复制本脚本到工作目录
  2. 修改 SOURCE / OUTPUT / PHOTO 路径
  3. python3 build_pages_style_resume.py
  4. osascript export_to_pages09.scpt output.docx output.pages

常量说明：
  - BASE_SIZE: 正文字号（pt），影响所有非标题文字
  - 顶部身份区、基本资料、职业意向的标题字号可分别调整
  - PHOTO 路径指向证件照；如不需要设为 ''
"""

import re
import zipfile
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Mm


# ============================================================
#  用户可配置常量（按需修改）
# ============================================================
SOURCE = Path("简历.md")          # 源 Markdown 文件
OUTPUT = Path("简历.docx")        # 输出的 Word 文件
PHOTO = Path("photo.png")         # 证件照路径，不存在时自动跳过

BASE_SIZE = 9.2                   # 正文字号（pt）
HEADER_FONT_SIZE = 9              # 顶部身份区字号
SECTION_TITLE_SIZE = 10           # 章节标题字号

FONT_CN = "PingFang SC"
FONT_LATIN = "Helvetica Neue"
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "D9D9D9"
BORDER_GRAY = "CFCFCF"
CONTENT_WIDTH = 10000             # 全宽列宽（dxa）


# ============================================================
#  字段映射表：中文简历常见表头 → 内部键名
#  添加新的映射即可支持更多字段名变体
# ============================================================
PROFILE_FIELD_ALIASES = {
    "姓名": ["姓名", "名字", "Name"],
    "性别": ["性别", "Sex"],
    "手机": ["手机", "电话", "手机号", "联系电话"],
    "邮箱": ["邮箱", "Email", "电子邮箱", "邮件"],
    "城市": ["城市", "所在地点", "所在地", "工作地点"],
    "工作年限": ["工作年限", "工作经验", "年限"],
    "求职意向": ["求职意向", "期望职位", "应聘职位", "职位"],
    "到岗时间": ["到岗时间", "可到岗时间", "入职时间"],
    "学历": ["学历", "最高学历"],
    "年龄": ["年龄"],
    "婚姻状况": ["婚姻状况", "婚姻"],
    "国籍": ["国籍"],
    "户籍": ["户籍", "户口"],
    "目前状态": ["目前状态", "当前状态"],
    "期望行业": ["期望行业", "行业"],
    "期望年薪": ["期望年薪", "期望薪资", "薪资", "期望薪酬"],
    "期望地点": ["期望地点"],
    "目前公司": ["目前公司", "当前公司"],
    "目前职位": ["目前职位", "当前职位"],
}


# ============================================================
#  OOXML 工具函数
# ============================================================

def set_run(run, size=BASE_SIZE, bold=False, color=INK):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color=BORDER_GRAY, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = tc_pr.first_child_found_in("w:tcBorders")
    if bdr is None:
        bdr = OxmlElement("w:tcBorders")
        tc_pr.append(bdr)
    for edge in ("top", "left", "bottom", "right"):
        el = bdr.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            bdr.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = tc_pr.first_child_found_in("w:tcBorders")
    if bdr is None:
        bdr = OxmlElement("w:tcBorders")
        tc_pr.append(bdr)
    for edge in ("top", "left", "bottom", "right"):
        el = bdr.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            bdr.append(el)
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")


def margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        el = tc_mar.find(qn(f"w:{key}"))
        if el is None:
            el = OxmlElement(f"w:{key}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def no_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(OxmlElement("w:noWrap"))


def widths(table, values):
    table.autofit = False
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for value in values:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(values[i] / 567)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)


def set_row_height(table, row_idx, height_pt):
    row = table.rows[row_idx]
    row.height = Pt(height_pt)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""


# ============================================================
#  单元格/段落写入辅助
# ============================================================

def cell_text(cell, text, size=BASE_SIZE, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def header_cell_text(cell, text, size=HEADER_FONT_SIZE, bold=False, color=INK):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell, top=30, start=60, bottom=30, end=60)
    cell_text(cell, text, size=size, bold=bold, color=color)


def add_p(container, text="", size=BASE_SIZE, bold=False, color=INK, before=0, after=2.2, indent=0):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    return p


def add_inline_label_p(container, label, text, size=BASE_SIZE):
    p = add_p(container, after=2.2)
    r = p.add_run(label)
    set_run(r, size=size, bold=True)
    r = p.add_run(text)
    set_run(r, size=size)


# ============================================================
#  Markdown 解析
# ============================================================

def parse_resume(md):
    """解析 Markdown 简历，返回 (profile_dict, sections_dict)"""
    lines = md.splitlines()
    profile = {}
    sections = {}
    current = None
    current_h3 = None

    for line in lines:
        s = line.strip()
        if s.startswith("|") and "项目" not in s and "---" not in s:
            cells = [c.strip() for c in s.strip("|").split("|")]
            if len(cells) == 2:
                profile[cells[0]] = cells[1]
        if s.startswith("## "):
            current = s[3:].strip()
            sections[current] = []
            current_h3 = None
        elif s.startswith("### ") and current:
            current_h3 = {"title": s[4:].strip(), "lines": []}
            sections[current].append(current_h3)
        elif current and s and not s.startswith("#"):
            target = current_h3["lines"] if current_h3 else sections[current]
            target.append(s)
    return profile, sections


def strip_md(text):
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text).strip()


def split_label(line):
    """将 '标签：值' 形式的行拆分为 (标签：, 值)"""
    clean = strip_md(line)
    m = re.match(r"([^：:]{2,12})[：:]\s*(.*)", clean)
    if m:
        return m.group(1) + "：", m.group(2)
    return "", clean


def get_profile_field(profile, key):
    """通过别名表查找 profile 字段值"""
    aliases = PROFILE_FIELD_ALIASES.get(key, [key])
    for alias in aliases:
        value = profile.get(alias)
        if value:
            return value
    return ""


def profile_value(profile, *keys, default=""):
    for key in keys:
        value = profile.get(key)
        if value:
            return value
    return default


# ============================================================
#  布局组件
# ============================================================

def add_section_title_row(doc, title):
    table = doc.add_table(rows=1, cols=1)
    widths(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    no_borders(cell)
    margins(cell, top=60, start=90, bottom=70, end=90)
    cell_text(cell, title, size=SECTION_TITLE_SIZE, bold=True)


def add_info_grid(doc, profile):
    """基本资料信息网格（5行4列）"""
    rows_data = [
        ("年　　龄：", get_profile_field(profile, "年龄"), "婚姻状况：", get_profile_field(profile, "婚姻状况")),
        ("手　　机：", get_profile_field(profile, "手机"), "邮　　箱：", get_profile_field(profile, "邮箱")),
        ("国　　籍：", get_profile_field(profile, "国籍"), "户　　籍：", get_profile_field(profile, "户籍")),
        ("目前状态：", get_profile_field(profile, "目前状态"), "所在地点：", get_profile_field(profile, "城市")),
        ("学　　历：", get_profile_field(profile, "学历"), "", ""),
    ]
    # 过滤全空行
    rows_data = [r for r in rows_data if any(v for v in r)]

    if not rows_data:
        return

    table = doc.add_table(rows=len(rows_data) + 1, cols=4)
    widths(table, [1250, 3300, 1550, 3900])

    title_cell = table.cell(0, 0).merge(table.cell(0, 3))
    no_borders(title_cell)
    margins(title_cell, top=50, start=90, bottom=80, end=90)
    cell_text(title_cell, "基本资料", size=SECTION_TITLE_SIZE, bold=True)

    for row, values in zip(table.rows[1:], rows_data):
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            no_borders(cell)
            is_label = idx in (0, 2)
            cell_text(cell, value, size=BASE_SIZE, bold=is_label, color=MUTED if is_label else INK)


def add_intention(doc, profile):
    """职业意向网格"""
    rows_data = [
        ("期望行业：", get_profile_field(profile, "期望行业")),
        ("期望职位：", get_profile_field(profile, "求职意向")),
        ("期望地点：", get_profile_field(profile, "期望地点") or get_profile_field(profile, "城市"),
         "期望年薪：", get_profile_field(profile, "期望年薪") or "面议"),
    ]
    # 过滤全空行
    rows_data = [r for r in rows_data if any(v for v in r)]
    if not rows_data:
        return

    table = doc.add_table(rows=len(rows_data) + 1, cols=4)
    widths(table, [1250, 5450, 1550, 1750])
    title_cell = table.cell(0, 0).merge(table.cell(0, 3))
    no_borders(title_cell)
    margins(title_cell, top=50, start=90, bottom=80, end=90)
    cell_text(title_cell, "职业意向", size=SECTION_TITLE_SIZE, bold=True)

    for r_idx, row in enumerate(table.rows[1:]):
        for c_idx, cell in enumerate(row.cells):
            no_borders(cell)
            value = rows_data[r_idx][c_idx] if c_idx < len(rows_data[r_idx]) else ""
            is_label = c_idx in (0, 2)
            cell_text(cell, value, size=BASE_SIZE, bold=is_label, color=MUTED if is_label else INK)


def add_gray_band(doc, left, middle, right=None, second=None):
    table = doc.add_table(rows=1, cols=3)
    widths(table, [1900, 5000, 3100])
    for cell in table.rows[0].cells:
        shade(cell, LIGHT_GRAY)
        borders(cell, color=LIGHT_GRAY, size="8")
    cell_text(table.cell(0, 0), left, size=BASE_SIZE, bold=True)
    cell_text(table.cell(0, 1), middle + (f"\n{second}" if second else ""), size=BASE_SIZE, bold=True)
    cell_text(table.cell(0, 2), right or "", size=BASE_SIZE, bold=True)


def add_label_block(doc, label, lines, numbered=True):
    """带编号的职责/成果列表"""
    if not lines:
        return
    table = doc.add_table(rows=1, cols=2)
    widths(table, [1350, 8650])
    for cell in table.rows[0].cells:
        no_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell_text(table.cell(0, 0), label, size=BASE_SIZE, bold=True)
    container = table.cell(0, 1)
    clear_paragraph(container.paragraphs[0])
    for idx, line in enumerate(lines, start=1):
        clean = strip_md(line[2:] if line.startswith("- ") else line)
        prefix = f"{idx}.　" if numbered else ""
        add_p(container, prefix + clean, size=BASE_SIZE, after=1.4)


def add_key_value_table(doc, rows, label_width=1700, value_width=8300):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    widths(table, [label_width, value_width])
    for row, (label, value) in zip(table.rows, rows):
        label_cell, value_cell = row.cells
        no_borders(value_cell)
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        shade(label_cell, LIGHT_GRAY)
        borders(label_cell, color=LIGHT_GRAY, size="8")
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell_text(label_cell, label, size=BASE_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_paragraph(value_cell.paragraphs[0])
        add_p(value_cell, value, size=BASE_SIZE, after=1.2)


def add_single_column_table(doc, lines):
    """纯文本段落列表（个人优势 / 自我评价）"""
    if not lines:
        return
    table = doc.add_table(rows=len(lines), cols=1)
    widths(table, [CONTENT_WIDTH])
    for row, text in zip(table.rows, lines):
        cell = row.cells[0]
        no_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        clear_paragraph(cell.paragraphs[0])
        add_p(cell, strip_md(text), size=BASE_SIZE, after=1.2)


def add_labeled_line_block(doc, label, text):
    if not text:
        return
    table = doc.add_table(rows=1, cols=2)
    widths(table, [1350, 8650])
    for cell in table.rows[0].cells:
        no_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell_text(table.cell(0, 0), label, size=BASE_SIZE, bold=True)
    clear_paragraph(table.cell(0, 1).paragraphs[0])
    add_p(table.cell(0, 1), text, size=BASE_SIZE, after=1.4)


# ============================================================
#  工作经历 / 项目经历 / 教育经历
# ============================================================

def add_experience(doc, sections, profile):
    for item in sections.get("工作经历", []):
        lines = item["lines"]
        role = ""
        dates = ""
        intro = []
        bullets = []
        in_bullets = False
        for line in lines:
            label, value = split_label(line)
            if label == "职位：":
                role = value
            elif label == "时间：":
                dates = value
            elif "主要工作内容" in line:
                in_bullets = True
            elif line.startswith("- ") or in_bullets:
                if line.startswith("- "):
                    bullets.append(line)
            else:
                intro.append(strip_md(line))
        location = get_profile_field(profile, "城市")
        add_gray_band(doc, dates, item["title"], f"工作地点：{location}" if location else "", f"职位：{role}")
        add_labeled_line_block(doc, "工作概况：", "".join(intro))
        add_label_block(doc, "工作职责：", bullets[:7])


def add_projects(doc, sections):
    for item in sections.get("项目经历", []):
        lines = item["lines"]
        meta = []
        desc = []
        duties = []
        results = []
        target = meta
        for line in lines:
            clean = strip_md(line)
            if "项目描述" in clean:
                target = desc
                continue
            if "主要职责" in clean:
                target = duties
                continue
            if "项目成果" in clean:
                target = results
                continue
            target.append(line)
        add_gray_band(doc, "", item["title"], "")
        meta_rows = []
        for line in meta[:5]:
            label, value = split_label(line)
            if label:
                meta_rows.append((label.rstrip("："), value))
        add_key_value_table(doc, meta_rows)
        if desc:
            add_labeled_line_block(doc, "项目描述：", "".join(strip_md(x) for x in desc))
        if duties:
            add_label_block(doc, "主要职责：", duties[:7])
        if results:
            add_label_block(doc, "项目成果：", results[:4])


def add_education(doc, sections):
    """
    支持两种教育经历格式：
      1. ### 学校名称（dict 格式）
      2. Markdown 表格（| 学校 | 学历 | 专业 | 时间 |）
    """
    items = sections.get("教育经历", [])
    if not items:
        return

    if isinstance(items[0], dict):
        # dict 格式：用灰色条展示
        for item in items:
            school = item["title"]
            degree = ""
            major = ""
            dates = ""
            for line in item["lines"]:
                label, value = split_label(line)
                if label in ("学历：", "层次："):
                    degree = value
                elif label == "专业：":
                    major = value
                elif label in ("时间：", "在校时间："):
                    dates = value
            add_gray_band(doc, dates, school, degree)
            if major:
                add_inline_label_p(doc, "专业：", major)
    else:
        # Markdown 表格格式：| 学校 | 学历 | 专业 | 时间 |
        rows = []
        in_header = True
        for line in items:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if len(cells) >= 4:
                if in_header and "学校" in cells[0]:
                    in_header = False
                    continue
                if "---" in cells[0]:
                    continue
                rows.append(cells[:4])
        if rows:
            tbl = doc.add_table(rows=len(rows), cols=1)
            widths(tbl, [CONTENT_WIDTH])
            for idx, (school, degree, major, dates) in enumerate(rows):
                cell = tbl.cell(idx, 0)
                margins(cell, top=50, start=100, bottom=50, end=100)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                parts = [major] if major and major != "专业" else []
                content = f"{dates}\u3000{school} | {degree}"
                if parts:
                    content += f" | {parts[0]}"
                cell_text(cell, content, size=BASE_SIZE, bold=False)


# ============================================================
#  空白页修复
# ============================================================

def fix_docx_blank_page(doc):
    """删除 docx 末尾空段落，防止 Pages 导入后多出空白页"""
    body = doc.element.body
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    children = list(body)
    for i in range(len(children) - 1, -1, -1):
        child = children[i]
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            texts = child.findall(f'{{{ns}}}t') if '}' in child.tag else []
            text = ''.join(t.text or '' for t in texts)
            if not text.strip():
                body.remove(child)
            else:
                break
        elif tag in ('tbl', 'sectPr'):
            break

    # 同时将最后一个段落的间距归零
    last_para = None
    for child in reversed(list(body)):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            last_para = child
            break
        elif tag == 'tbl':
            paras = child.findall(f'.//{{{ns}}}p')
            if paras:
                last_para = paras[-1]
            break

    if last_para is not None:
        pPr = last_para.find(f'{{{ns}}}pPr')
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            last_para.insert(0, pPr)
        spacing = pPr.find(f'{{{ns}}}spacing')
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')


def fix_pages_blank_page(pages_path):
    """修复 .pages 包中末尾空白 section 导致的空白页"""
    tmp_path = pages_path + '.tmp'
    try:
        with zipfile.ZipFile(pages_path, 'r') as zin:
            content = zin.read('index.xml')
        text = content.decode('utf-8', errors='replace')
        idx = text.find('sf:section-1')
        if idx >= 0:
            section_start = text.rfind('<sf:section', 0, idx)
            section_end = text.find('</sf:section>', idx) + len('</sf:section>')
            if section_start >= 0 and section_end > section_start:
                new_text = text[:section_start] + text[section_end:]
                with zipfile.ZipFile(pages_path, 'r') as zin2:
                    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                        for item in zin2.infolist():
                            if item.filename == 'index.xml':
                                zout.writestr(item, new_text.encode('utf-8'))
                            else:
                                zout.writestr(item, zin2.read(item.filename))
                os.replace(tmp_path, pages_path)
                return True
        return False
    except Exception:
        return False


# ============================================================
#  构建
# ============================================================

def build():
    profile, sections = parse_resume(SOURCE.read_text(encoding="utf-8"))
    doc = Document()

    # A4 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # 默认样式
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(BASE_SIZE)
    normal.paragraph_format.space_after = Pt(1)

    # ---------- 顶部身份区（3行×5列） ----------
    header = doc.add_table(rows=3, cols=5)
    widths(header, [1627, 1803, 2615, 1648, 2339])
    set_row_height(header, 0, 24)
    set_row_height(header, 1, 26.65)
    set_row_height(header, 2, 37.05)
    for row in header.rows:
        for cell in row.cells:
            shade(cell, LIGHT_GRAY)
            borders(cell, color=LIGHT_GRAY, size="8")

    # 照片格（A1:A3 合并）
    photo_cell = header.cell(0, 0).merge(header.cell(2, 0))
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(photo_cell, top=60, start=25, bottom=0, end=25)
    clear_paragraph(photo_cell.paragraphs[0])
    photo_p = photo_cell.paragraphs[0]
    photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_p.paragraph_format.space_before = Pt(0)
    photo_p.paragraph_format.space_after = Pt(0)
    photo_p.paragraph_format.line_spacing = 1.0
    if PHOTO and PHOTO.exists():
        photo_run = photo_p.add_run()
        photo_run.add_picture(str(PHOTO), width=Cm(2.45), height=Cm(3.27))

    # 从简历推断当前公司和职位
    current_company = get_profile_field(profile, "目前公司")
    if not current_company and sections.get("工作经历"):
        first = sections["工作经历"][0]
        if isinstance(first, dict):
            current_company = first["title"]

    current_title = get_profile_field(profile, "目前职位") or get_profile_field(profile, "求职意向")
    current_title = current_title.split("/")[0].strip() if current_title else ""

    name = get_profile_field(profile, "姓名") or "姓名"
    gender = get_profile_field(profile, "性别")
    work_years = get_profile_field(profile, "工作年限")

    header_cell_text(header.cell(0, 1), name, size=15, bold=True)
    header_cell_text(header.cell(0, 2), gender, size=HEADER_FONT_SIZE)
    header_cell_text(header.cell(0, 3), "", size=HEADER_FONT_SIZE)
    header_cell_text(header.cell(0, 4), "", size=HEADER_FONT_SIZE)
    header_cell_text(header.cell(1, 1), "目前公司：", size=HEADER_FONT_SIZE, color=MUTED)
    no_wrap(header.cell(1, 2))
    header_cell_text(header.cell(1, 2), current_company, size=HEADER_FONT_SIZE, bold=True)
    header_cell_text(header.cell(1, 3), "", size=HEADER_FONT_SIZE)
    header_cell_text(header.cell(1, 4), "", size=HEADER_FONT_SIZE)
    header_cell_text(header.cell(2, 1), "目前职位：", size=HEADER_FONT_SIZE, color=MUTED)
    header_cell_text(header.cell(2, 2), current_title, size=HEADER_FONT_SIZE, bold=True)
    header_cell_text(header.cell(2, 3), "工作年限：", size=HEADER_FONT_SIZE, bold=True)
    header_cell_text(header.cell(2, 4), work_years, size=HEADER_FONT_SIZE, bold=True)

    # ---------- 内容区 ----------
    add_info_grid(doc, profile)
    add_intention(doc, profile)

    for section_name in ("个人优势", "专业技能", "工作经历", "项目经历", "教育经历", "自我评价"):
        if section_name not in sections or not sections[section_name]:
            continue
        add_section_title_row(doc, section_name)
        if section_name == "个人优势":
            add_single_column_table(doc, sections["个人优势"])
        elif section_name == "专业技能":
            skill_rows = []
            for line in sections["专业技能"]:
                label, value = split_label(line)
                if label:
                    skill_rows.append((label.rstrip("："), value))
            add_key_value_table(doc, skill_rows)
        elif section_name == "工作经历":
            add_experience(doc, sections, profile)
        elif section_name == "项目经历":
            add_projects(doc, sections)
        elif section_name == "教育经历":
            add_education(doc, sections)
        elif section_name == "自我评价":
            add_single_column_table(doc, sections["自我评价"])

    # 修复空白页
    fix_docx_blank_page(doc)
    doc.save(OUTPUT)
    print(f"✅ 已生成 {OUTPUT}")


def fix_pages(pages_path):
    """对已生成的 .pages 文件修复空白页"""
    pages = Path(pages_path)
    if pages.exists():
        if fix_pages_blank_page(str(pages)):
            print(f"✅ 已修复 {pages.name} 空白页")
            return True
        else:
            print(f"ℹ️  {pages.name} 无需修复")
            return False
    else:
        print(f"❌ 文件不存在: {pages_path}")
        return False


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "fix-pages":
        # 单独修复 .pages 空白页：python3 script.py fix-pages output.pages
        fix_pages(sys.argv[2] if len(sys.argv) > 2 else OUTPUT.with_suffix('.pages'))
    else:
        build()
        print("\n💡 导出 Pages 后别忘了修复空白页：")
        print(f"   python3 {__import__('os').path.basename(__file__)} fix-pages {OUTPUT.with_suffix('.pages').name}")
